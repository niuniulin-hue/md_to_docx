"""
Walk a mistune AST and build a python-docx Document.

Supported GFM elements
-----------------------
Block:  heading, paragraph, blank_line, thematic_break (hr),
        block_code (fenced + indented), block_quote,
        list (ordered / unordered / task), table
Inline: text, softbreak, linebreak, codespan, strong, emphasis,
        strikethrough, link, image, raw_html (stripped)
"""

from __future__ import annotations

import io
import re
import unicodedata
from typing import Mapping, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .styles import (
    apply_styles,
    apply_publish_formatting,
    add_paragraph_shading,
    add_left_border,
    add_horizontal_rule,
    set_rfonts_on_rpr,
    set_run_fonts,
)


# ---------------------------------------------------------------------------
# Public entry-point
# ---------------------------------------------------------------------------

KDP_SAFE_ICON_MAP: dict[str, str] = {
    # Status / callouts
    "✅": "[OK]",
    "✔": "[OK]",
    "☑": "[x]",
    "☐": "[ ]",
    "❌": "[X]",
    "✖": "[X]",
    "⚠": "[Warning]",
    "🚨": "[Alert]",
    "🚫": "[No]",
    "🛑": "[Stop]",
    "🆘": "[Help]",
    "‼": "!!",
    "❗": "!",
    "❕": "!",
    "ℹ": "[Info]",
    "❓": "[Question]",
    "❔": "[Question]",
    "⁉": "[Question]",

    # Notes / references / structure
    "💡": "[Tip]",
    "🎯": "[Goal]",
    "📌": "[Pin]",
    "📝": "[Note]",
    "📘": "[Book]",
    "📖": "[Book]",
    "📋": "[List]",
    "📦": "[Package]",
    "📅": "[Calendar]",
    "🗓️": "[Calendar]",
    "📜": "[Guide]",
    "🔍": "[Review]",
    "🔗": "[Link]",
    "🔒": "[Lock]",
    "🔓": "[Open]",

    # Progress / charts / navigation
    "🔥": "[Hot]",
    "🚀": "[Launch]",
    "🏁": "[Finish]",
    "🔄": "[Refresh]",
    "📊": "[Chart]",
    "📈": "[Up]",
    "📉": "[Down]",
    "📏": "[Measure]",
    "⚖️": "[Balance]",
    "⚓": "[Anchor]",
    "🧭": "[Direction]",
    "🗺️": "[Map]",
    "⭐": "*",
    "🌟": "*",
    "🎉": "[Success]",
    "🏆": "[Award]",
    "👣": "[Steps]",
    "➡": "->",
    "➜": "->",
    "➔": "->",
    "➤": "->",
    "👉": "->",
    "←": "<-",
    "⬅": "<-",
    "↔": "<->",
    "↩": "<-",
    "↪": "->",
    "↑": "^",
    "⬆": "^",
    "↓": "v",
    "⬇": "v",

    # Time / routine
    "⏰": "[Time]",
    "⏱️": "[Time]",
    "🕒": "[Time]",
    "🕰️": "[Time]",
    "☀️": "[Day]",
    "🌞": "[Day]",
    "🌤️": "[Day]",
    "🌅": "[Morning]",
    "🌇": "[Evening]",
    "🌆": "[Evening]",
    "🌙": "[Night]",
    "🌜": "[Night]",
    "🌛": "[Night]",
    "🌃": "[Night]",
    "🌑": "[Night]",

    # Health / activity / work
    "💪": "[Strength]",
    "🏃": "[Run]",
    "🏃‍♀️": "[Run]",
    "🚶": "[Walk]",
    "🚶‍♀️": "[Walk]",
    "🧘": "[Calm]",
    "🧘‍♀️": "[Calm]",
    "🏋️‍♀️": "[Strength]",
    "🏋️‍♂️": "[Strength]",
    "🏊‍♀️": "[Swim]",
    "🚴": "[Ride]",
    "🚴‍♀️": "[Ride]",
    "🤸‍♀️": "[Stretch]",
    "😴": "[Sleep]",
    "💤": "[Sleep]",
    "🛌": "[Rest]",
    "🩺": "[Health]",
    "🏥": "[Health]",
    "💊": "[Medication]",
    "🩸": "[Health]",
    "💧": "[Water]",
    "💼": "[Work]",
    "💻": "[Laptop]",
    "📱": "[Phone]",
    "🏠": "[Home]",

    # Food / drink
    "🍽️": "[Meal]",
    "🍱": "[Meal]",
    "🥗": "[Food]",
    "🥣": "[Food]",
    "🍎": "[Fruit]",
    "🥬": "[Vegetable]",
    "🥦": "[Vegetable]",
    "🍚": "[Grain]",
    "🍗": "[Protein]",
    "🐟": "[Protein]",
    "🥛": "[Drink]",
    "☕": "[Drink]",
    "🍵": "[Drink]",

    # General tone / emphasis
    "❤️": "[Heart]",
    "💖": "[Heart]",
    "💕": "[Heart]",
}

_TEXT_NORMALIZATION_TABLE = str.maketrans({
    "\u200d": None,  # zero-width joiner
    "\ufe0e": None,  # text presentation selector
    "\ufe0f": None,  # emoji presentation selector
})

_ZERO_WIDTH_JOINER = "\u200d"
_COMBINING_KEYCAP = "\u20e3"
_SKIN_TONE_MODIFIER_RANGE = range(0x1F3FB, 0x1F400)
_REGIONAL_INDICATOR_RANGE = range(0x1F1E6, 0x1F200)
_ICON_NAME_STOPWORDS = {
    "and",
    "button",
    "digit",
    "for",
    "letter",
    "of",
    "selector",
    "sign",
    "symbol",
    "squared",
    "variation",
    "with",
}
_ICON_NAME_REWRITES = {
    "aesculapius": "health",
    "ball of yarn": "yarn",
    "clockwise": "refresh",
    "counterclockwise": "refresh",
    "droplet": "water",
    "laptop computer": "laptop",
    "mobile phone": "phone",
    "person in lotus position": "meditation",
    "running": "run",
    "runner": "run",
    "sleeping": "sleep",
    "stethoscope": "health",
    "sunrise": "morning",
    "sunset": "evening",
    "walking": "walk",
    "weight lifter": "strength",
}


def _is_variation_selector(ch: str) -> bool:
    return ord(ch) in {0xFE0E, 0xFE0F}


def _is_skin_tone_modifier(ch: str) -> bool:
    return ord(ch) in _SKIN_TONE_MODIFIER_RANGE


def _is_regional_indicator(ch: str) -> bool:
    return ord(ch) in _REGIONAL_INDICATOR_RANGE


def _is_icon_like_char(ch: str) -> bool:
    cp = ord(ch)
    category = unicodedata.category(ch)
    return (
        ch == _ZERO_WIDTH_JOINER
        or ch == _COMBINING_KEYCAP
        or _is_variation_selector(ch)
        or _is_skin_tone_modifier(ch)
        or _is_regional_indicator(ch)
        or 0x2100 <= cp <= 0x214F
        or 0x2190 <= cp <= 0x21FF
        or 0x2300 <= cp <= 0x23FF
        or 0x2460 <= cp <= 0x24FF
        or 0x2600 <= cp <= 0x27BF
        or 0x2900 <= cp <= 0x2BFF
        or 0x1F000 <= cp <= 0x1FAFF
        or category == "So"
    )


def _starts_keycap_cluster(text: str, start: int) -> bool:
    if text[start] not in "0123456789#*":
        return False
    next_index = start + 1
    if next_index < len(text) and _is_variation_selector(text[next_index]):
        next_index += 1
    return next_index < len(text) and text[next_index] == _COMBINING_KEYCAP


def _consume_icon_cluster(text: str, start: int) -> tuple[str, int]:
    first = text[start]
    if _starts_keycap_cluster(text, start):
        end = start + 1
        if end < len(text) and _is_variation_selector(text[end]):
            end += 1
        if end < len(text) and text[end] == _COMBINING_KEYCAP:
            end += 1
        return text[start:end], end
    if _is_regional_indicator(first):
        end = start + 1
        if end < len(text) and _is_regional_indicator(text[end]):
            end += 1
        return text[start:end], end

    end = start + 1
    while end < len(text):
        ch = text[end]
        prev = text[end - 1]
        if ch == _COMBINING_KEYCAP or _is_variation_selector(ch) or _is_skin_tone_modifier(ch):
            end += 1
            continue
        if ch == _ZERO_WIDTH_JOINER:
            end += 1
            continue
        if prev == _ZERO_WIDTH_JOINER and _is_icon_like_char(ch):
            end += 1
            continue
        break
    return text[start:end], end


def _title_case_label(words: list[str]) -> str:
    return " ".join(word.capitalize() for word in words if word)


def _fallback_label_for_cluster(cluster: str) -> str | None:
    normalized = cluster.translate(_TEXT_NORMALIZATION_TABLE)
    if not normalized:
        return None

    if all(_is_regional_indicator(ch) for ch in normalized):
        country_code = "".join(chr(ord("A") + ord(ch) - 0x1F1E6) for ch in normalized)
        return f"[Flag {country_code}]"

    if _COMBINING_KEYCAP in cluster:
        base = next((ch for ch in cluster if ch.isdigit() or ch in "#*"), None)
        if base:
            return f"[{base}]"

    name_parts: list[str] = []
    for ch in normalized:
        if ch == _COMBINING_KEYCAP:
            continue
        name = unicodedata.name(ch, "")
        if not name:
            continue
        lowered = name.lower().replace("-", " ")
        lowered = lowered.replace("heavy ", "")
        lowered = lowered.replace("black ", "")
        lowered = lowered.replace("white ", "")
        lowered = lowered.replace("medium ", "")
        lowered = lowered.replace("small ", "")
        lowered = lowered.replace("large ", "")
        for source, target in _ICON_NAME_REWRITES.items():
            lowered = lowered.replace(source, target)
        for part in lowered.split():
            if part and part not in _ICON_NAME_STOPWORDS and part not in name_parts:
                name_parts.append(part)

    if not name_parts:
        return "[Symbol]"

    compact = _title_case_label(name_parts[:4])
    return f"[{compact}]" if compact else "[Symbol]"


def build_text_replacements(
    *,
    kdp_safe_icons: bool = False,
    icon_map: Mapping[str, str] | None = None,
) -> tuple[tuple[str, str], ...]:
    """Build an ordered replacement list for prose text written to DOCX."""
    merged: dict[str, str] = {}
    if kdp_safe_icons:
        for source, target in KDP_SAFE_ICON_MAP.items():
            normalized_source = source.translate(_TEXT_NORMALIZATION_TABLE)
            if normalized_source:
                merged[normalized_source] = target
    if icon_map:
        for source, target in icon_map.items():
            normalized_source = source.translate(_TEXT_NORMALIZATION_TABLE)
            if normalized_source:
                merged[normalized_source] = target
    return tuple(sorted(merged.items(), key=lambda item: len(item[0]), reverse=True))

def ast_to_docx(
    tokens: list,
    doc: Optional[DocxDocument] = None,
    *,
    kdp_safe_icons: bool = False,
    icon_map: Mapping[str, str] | None = None,
    publish_mode: bool = False,
) -> DocxDocument:
    """Convert a mistune token list to a python-docx Document."""
    if doc is None:
        doc = Document()
    apply_styles(doc)
    if publish_mode:
        apply_publish_formatting(doc)
    ctx = _Context(
        doc,
        text_replacements=build_text_replacements(
            kdp_safe_icons=kdp_safe_icons,
            icon_map=icon_map,
        ),
        auto_label_icons=kdp_safe_icons,
    )
    _render_tokens(tokens, ctx)
    return doc


# ---------------------------------------------------------------------------
# Internal context object
# ---------------------------------------------------------------------------

class _Context:
    def __init__(
        self,
        doc: DocxDocument,
        text_replacements: tuple[tuple[str, str], ...] = (),
        auto_label_icons: bool = False,
    ):
        self.doc = doc
        self.text_replacements = text_replacements
        self.auto_label_icons = auto_label_icons
        self.list_depth = 0          # current nesting depth (0 = top-level)
        self.list_ordered_stack: list[bool] = []  # True = ordered at each level
        self.list_item_index: list[int] = []       # counter per level


def _normalize_text(text: str, ctx: _Context, is_code: bool = False) -> str:
    if not text:
        return text
    if is_code and not ctx.auto_label_icons:
        return text
    if not ctx.text_replacements and not ctx.auto_label_icons:
        return text

    replacement_lookup = dict(ctx.text_replacements)
    parts: list[str] = []
    index = 0

    while index < len(text):
        if not _is_icon_like_char(text[index]) and not _starts_keycap_cluster(text, index):
            parts.append(text[index])
            index += 1
            continue

        cluster, next_index = _consume_icon_cluster(text, index)
        normalized_cluster = cluster.translate(_TEXT_NORMALIZATION_TABLE)

        replacement = replacement_lookup.get(normalized_cluster)
        if replacement is not None:
            parts.append(replacement)
        elif ctx.auto_label_icons and normalized_cluster:
            parts.append(_fallback_label_for_cluster(cluster) or normalized_cluster)
        else:
            parts.append(cluster)

        index = next_index

    return "".join(parts)


# ---------------------------------------------------------------------------
# Block renderer dispatcher
# ---------------------------------------------------------------------------

def _render_tokens(tokens: list, ctx: _Context):
    for token in tokens:
        _render_block(token, ctx)


def _render_block(token: dict, ctx: _Context):
    t = token.get("type", "")

    if t == "heading":
        _render_heading(token, ctx)
    elif t in ("paragraph", "block_text"):
        _render_paragraph(token, ctx)
    elif t in ("blank_line",):
        pass  # skip blank lines
    elif t == "thematic_break":
        add_horizontal_rule(ctx.doc)
    elif t == "block_code":
        _render_block_code(token, ctx)
    elif t == "block_quote":
        _render_block_quote(token, ctx)
    elif t == "list":
        _render_list(token, ctx)
    elif t == "table":
        _render_table(token, ctx)
    elif t == "html":
        # raw HTML blocks — ignore tags, try to render text content
        text = re.sub(r"<[^>]+>", "", token.get("raw", "")).strip()
        if text:
            p = ctx.doc.add_paragraph()
            _add_run(p, text, ctx)
    else:
        # Fallback: try children
        children = token.get("children")
        if children:
            _render_tokens(children, ctx)


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _render_heading(token: dict, ctx: _Context):
    level = token.get("attrs", {}).get("level", 1)
    level = max(1, min(6, int(level)))
    p = ctx.doc.add_heading("", level=level)
    _fill_inline(p, token.get("children", []), ctx)


def _render_paragraph(token: dict, ctx: _Context):
    p = ctx.doc.add_paragraph()
    _fill_inline(p, token.get("children", []), ctx)


def _render_block_code(token: dict, ctx: _Context):
    raw = token.get("raw", "")
    # Split into lines to preserve whitespace
    for line in raw.splitlines():
        safe_line = _normalize_text(line, ctx, is_code=True)
        p = ctx.doc.add_paragraph(style="Code Block")
        run = p.add_run(safe_line)
        set_run_fonts(run, safe_line, is_code=True)
        run.font.size = Pt(9)
        add_paragraph_shading(p, "F6F8FA")


def _render_block_quote(token: dict, ctx: _Context):
    children = token.get("children", [])
    for child in children:
        child_type = child.get("type", "")
        if child_type in ("paragraph", "block_text"):
            p = ctx.doc.add_paragraph(style="Block Quote")
            add_left_border(p, "0969DA", width_eighths=18)
            _fill_inline(p, child.get("children", []), ctx)
        elif child_type == "block_quote":
            # Nested blockquote — recurse
            _render_block_quote(child, ctx)
        else:
            _render_block(child, ctx)


def _render_list(token: dict, ctx: _Context):
    ordered = token.get("attrs", {}).get("ordered", False)
    ctx.list_depth += 1
    ctx.list_ordered_stack.append(ordered)
    ctx.list_item_index.append(0)

    depth = ctx.list_depth  # capture for style lookup
    for item in token.get("children", []):
        item_type = item.get("type", "")
        if item_type == "task_list_item":
            _render_task_list_item(item, depth, ctx)
        else:
            _render_list_item(item, ordered, depth, ctx)

    ctx.list_depth -= 1
    ctx.list_ordered_stack.pop()
    ctx.list_item_index.pop()


def _render_list_item(token: dict, ordered: bool, depth: int, ctx: _Context):
    """Render a single list_item token."""
    children = token.get("children", [])
    paragraph_added = False
    p = None

    for child in children:
        child_type = child.get("type", "")

        if child_type == "list":
            # Nested list — recurse
            _render_list(child, ctx)
        elif child_type in ("paragraph", "block_text"):
            p = _add_list_paragraph(ctx.doc, ordered, depth)
            paragraph_added = True
            _fill_inline(p, child.get("children", []), ctx)
        else:
            # Inline content directly
            if not paragraph_added or p is None:
                p = _add_list_paragraph(ctx.doc, ordered, depth)
                paragraph_added = True
            _fill_inline(p, child.get("children", []) or [], ctx)


def _render_task_list_item(token: dict, depth: int, ctx: _Context):
    """Render a task_list_item token (checkbox item)."""
    checked = token.get("attrs", {}).get("checked", False)
    checkbox = "☑ " if checked else "☐ "
    children = token.get("children", [])

    paragraph_added = False
    p = None
    for child in children:
        child_type = child.get("type", "")
        if child_type == "list":
            _render_list(child, ctx)
        elif child_type in ("paragraph", "block_text"):
            p = _add_list_paragraph(ctx.doc, False, depth)
            paragraph_added = True
            _add_run(p, checkbox, ctx)
            _fill_inline(p, child.get("children", []), ctx)
        else:
            if not paragraph_added or p is None:
                p = _add_list_paragraph(ctx.doc, False, depth)
                _add_run(p, checkbox, ctx)
                paragraph_added = True
            _fill_inline(p, child.get("children", []) or [], ctx)


def _add_list_paragraph(doc: DocxDocument, ordered: bool, depth: int):
    """Create a paragraph with appropriate list style and indentation."""
    if depth > 1:
        style_name = f"List {'Number' if ordered else 'Bullet'} {min(depth, 3)}"
    else:
        style_name = f"List {'Number' if ordered else 'Bullet'}"
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph(style="List Bullet")
    # Extra indent for deep nesting
    if depth > 3:
        p.paragraph_format.left_indent = Cm(depth * 0.6)
    return p


def _render_table(token: dict, ctx: _Context):
    """Render a table token. mistune structure:
       table -> [table_head, table_body]
       table_head -> [table_cell, ...]   (head=True)
       table_body -> [table_row, ...]
       table_row  -> [table_cell, ...]
    """
    children = token.get("children", [])
    head_token = next((c for c in children if c.get("type") == "table_head"), None)
    body_token = next((c for c in children if c.get("type") == "table_body"), None)

    header_cells = head_token.get("children", []) if head_token else []
    body_rows = body_token.get("children", []) if body_token else []

    num_cols = len(header_cells)
    if num_cols == 0:
        return

    total_rows = 1 + len(body_rows)
    table = ctx.doc.add_table(rows=total_rows, cols=num_cols)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for ci, cell_token in enumerate(header_cells):
        cell = hrow.cells[ci]
        p = cell.paragraphs[0]
        _fill_inline(p, cell_token.get("children", []), ctx)
        for run in p.runs:
            run.bold = True
        _shade_table_cell(cell, "F6F8FA")

    # Body rows
    for ri, row_token in enumerate(body_rows):
        row = table.rows[ri + 1]
        row_cells = row_token.get("children", [])
        for ci, cell_token in enumerate(row_cells[:num_cols]):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            _fill_inline(p, cell_token.get("children", []), ctx)


def _shade_table_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Inline renderer
# ---------------------------------------------------------------------------

def _fill_inline(paragraph, tokens: list, ctx: _Context):
    """Add inline runs to *paragraph* from a list of inline tokens."""
    for token in tokens:
        _render_inline(paragraph, token, ctx)


def _render_inline(paragraph, token: dict, ctx: _Context):
    t = token.get("type", "")

    if t == "text":
        _add_run(paragraph, token.get("raw", ""), ctx)

    elif t in ("softbreak", "softline_break"):
        _add_run(paragraph, " ", ctx)

    elif t == "linebreak":
        _add_run(paragraph, "\n", ctx)

    elif t == "codespan":
        code_text = _normalize_text(token.get("raw", ""), ctx, is_code=True)
        run = paragraph.add_run(code_text)
        set_run_fonts(run, code_text, is_code=True)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xE3, 0x11, 0x6C)

    elif t == "strong":
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, bold=True)

    elif t == "emphasis":
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, italic=True)

    elif t == "strikethrough":
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, strike=True)

    elif t == "link":
        href = token.get("attrs", {}).get("url", "")
        _add_hyperlink(paragraph, token.get("children", []), href, ctx)

    elif t == "image":
        attrs = token.get("attrs", {})
        src = attrs.get("url", "")
        alt = attrs.get("alt", "")
        _add_image(paragraph, src, alt, ctx)

    elif t in ("html", "raw_html"):
        raw = token.get("raw", "")
        text = re.sub(r"<[^>]+>", "", raw)
        if text:
            _add_run(paragraph, text, ctx)

    else:
        # Recurse into children for unknown inline tokens
        for child in token.get("children", []):
            _render_inline(paragraph, child, ctx)


def _render_inline_with_style(paragraph, token: dict, ctx: _Context,
                               bold=False, italic=False, strike=False):
    t = token.get("type", "")

    if t == "text":
        run = _add_run(paragraph, token.get("raw", ""), ctx)
        if run is not None:
            run.bold = bold or None
            run.italic = italic or None
            if strike:
                run.font.strike = True

    elif t == "strong":
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, bold=True, italic=italic, strike=strike)

    elif t == "emphasis":
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, bold=bold, italic=True, strike=strike)

    elif t == "strikethrough":
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, bold=bold, italic=italic, strike=True)

    elif t == "codespan":
        code_text = _normalize_text(token.get("raw", ""), ctx, is_code=True)
        run = paragraph.add_run(code_text)
        set_run_fonts(run, code_text, is_code=True)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xE3, 0x11, 0x6C)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True

    elif t == "link":
        href = token.get("attrs", {}).get("url", "")
        _add_hyperlink(paragraph, token.get("children", []), href, ctx, bold=bold, italic=italic, strike=strike)

    elif t == "image":
        attrs = token.get("attrs", {})
        _add_image(paragraph, attrs.get("url", ""), attrs.get("alt", ""), ctx)

    else:
        for child in token.get("children", []):
            _render_inline_with_style(paragraph, child, ctx, bold=bold, italic=italic, strike=strike)


def _add_run(paragraph, text: str, ctx: _Context | None = None, is_code: bool = False):
    if not text:
        return None
    if ctx is not None:
        text = _normalize_text(text, ctx, is_code=is_code)
        if not text:
            return None
    run = paragraph.add_run(text)
    set_run_fonts(run, text, is_code=is_code)
    return run


# ---------------------------------------------------------------------------
# Hyperlinks
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, children: list, href: str, ctx: _Context, bold=False, italic=False, strike=False):
    """Insert a clickable hyperlink into *paragraph*."""
    part = paragraph.part
    r_id = part.relate_to(
        href,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    link_text = _normalize_text(_extract_text(children), ctx)
    if not link_text:
        link_text = href

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1")
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(color_el)
    rPr.append(u_el)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if strike:
        rPr.append(OxmlElement("w:strike"))
    set_rfonts_on_rpr(rPr, text=link_text)
    r.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = link_text
    t_el.set(qn("xml:space"), "preserve")
    r.append(t_el)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _extract_text(tokens: list) -> str:
    """Recursively extract plain text from an inline token list."""
    parts = []
    for t in tokens:
        token_type = t.get("type")
        if token_type == "text":
            parts.append(t.get("raw", ""))
        elif token_type == "codespan":
            parts.append(t.get("raw", ""))
        elif token_type in ("html", "raw_html"):
            parts.append(re.sub(r"<[^>]+>", "", t.get("raw", "")))
        elif token_type in ("softbreak", "softline_break", "linebreak"):
            parts.append(" ")
        else:
            parts.append(_extract_text(t.get("children", [])))
    return "".join(parts)


# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------

def _add_image(paragraph, src: str, alt: str, ctx: _Context):
    """Embed an image or fall back to alt text."""
    try:
        img_data = _load_image(src)
        if img_data:
            run = paragraph.add_run()
            run.add_picture(img_data, width=Inches(4.5))
            return
    except Exception:
        pass
    fallback_text = f"[Image: {alt}]" if alt else (f"[Image: {src}]" if src else "")
    if fallback_text:
        run = _add_run(paragraph, fallback_text, ctx)
        if run is not None:
            run.italic = True


def _load_image(src: str):
    """Return a BytesIO for the image at *src* (local path or URL)."""
    import os

    if src.startswith("http://") or src.startswith("https://"):
        import requests
        resp = requests.get(src, timeout=10)
        resp.raise_for_status()
        return io.BytesIO(resp.content)
    if os.path.isfile(src):
        return open(src, "rb")
    return None

