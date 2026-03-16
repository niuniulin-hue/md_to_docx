"""
Define and apply Word styles for the converted document.
"""
from typing import TypedDict
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULT_LATIN_FONT = "Calibri"
DEFAULT_CJK_FONT = "Microsoft JhengHei"
DEFAULT_JAPANESE_FONT = "Yu Gothic"
DEFAULT_CODE_FONT = "Consolas"
DEFAULT_CODE_CJK_FONT = "MS Gothic"


def apply_styles(doc: Document) -> None:
    """Add/configure all custom styles used by the renderer."""
    _ensure_base_document_fonts(doc)
    _ensure_code_block_style(doc)
    _ensure_inline_code_style(doc)
    _ensure_quote_style(doc)
    _ensure_list_styles(doc)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_or_create_style(doc, name, base_name=None, style_type="paragraph"):
    from docx.enum.style import WD_STYLE_TYPE
    type_map = {
        "paragraph": WD_STYLE_TYPE.PARAGRAPH,
        "character": WD_STYLE_TYPE.CHARACTER,
    }
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, type_map[style_type])
        if base_name:
            try:
                style.base_style = doc.styles[base_name]
            except KeyError:
                pass
        return style


def _contains_japanese(text: str) -> bool:
    return any(
        "\u3040" <= char <= "\u30ff"
        or "\u31f0" <= char <= "\u31ff"
        or "\uff66" <= char <= "\uff9d"
        for char in text
    )


def _contains_cjk(text: str) -> bool:
    return any(
        "\u3400" <= char <= "\u4dbf"
        or "\u4e00" <= char <= "\u9fff"
        or "\u3000" <= char <= "\u303f"
        or "\u3100" <= char <= "\u312f"
        or "\uff00" <= char <= "\uffef"
        for char in text
    )


def resolve_font_family(text: str, is_code: bool = False) -> dict[str, str]:
    """Return the best Latin/East Asia font pair for the given text."""
    latin_font = DEFAULT_CODE_FONT if is_code else DEFAULT_LATIN_FONT
    east_asia_font = DEFAULT_CODE_CJK_FONT if is_code else DEFAULT_CJK_FONT
    if _contains_japanese(text):
        east_asia_font = DEFAULT_CODE_CJK_FONT if is_code else DEFAULT_JAPANESE_FONT
    elif _contains_cjk(text):
        east_asia_font = DEFAULT_CODE_CJK_FONT if is_code else DEFAULT_CJK_FONT
    return {
        "ascii": latin_font,
        "hAnsi": latin_font,
        "cs": latin_font,
        "eastAsia": east_asia_font,
    }


def _get_or_add_xml_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def set_rfonts_on_rpr(rpr, text: str = "", is_code: bool = False):
    """Apply OOXML font mapping to an rPr element for multilingual text."""
    font_map = resolve_font_family(text, is_code=is_code)
    r_fonts = _get_or_add_xml_child(rpr, "w:rFonts")
    for attr_name, font_name in font_map.items():
        r_fonts.set(qn(f"w:{attr_name}"), font_name)


def set_run_fonts(run, text: str = "", is_code: bool = False):
    """Apply Latin + East Asia font settings to a python-docx run."""
    font_map = resolve_font_family(text, is_code=is_code)
    run.font.name = font_map["ascii"]
    set_rfonts_on_rpr(run._element.get_or_add_rPr(), text=text, is_code=is_code)


def _configure_style_fonts(style, text_sample: str = "", is_code: bool = False, size: Pt | None = None):
    font_map = resolve_font_family(text_sample, is_code=is_code)
    style.font.name = font_map["ascii"]
    if size is not None:
        style.font.size = size
    set_rfonts_on_rpr(_get_or_add_xml_child(style.element, "w:rPr"), text=text_sample, is_code=is_code)


def _ensure_base_document_fonts(doc):
    _configure_style_fonts(doc.styles["Normal"], text_sample="繁體中文 日本語 Français Español Deutsch")
    for heading_name in [f"Heading {level}" for level in range(1, 7)]:
        try:
            _configure_style_fonts(doc.styles[heading_name], text_sample="繁體中文 日本語 Français Español Deutsch")
        except KeyError:
            continue


def _ensure_code_block_style(doc):
    style = _get_or_create_style(doc, "Code Block", "Normal")
    _configure_style_fonts(style, text_sample="程式碼 コード code", is_code=True, size=Pt(9))
    font = style.font
    font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    pf = style.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.5)
    _set_paragraph_shading(style.element, "F6F8FA")


def _ensure_inline_code_style(doc):
    style = _get_or_create_style(doc, "Inline Code", "Default Paragraph Font", style_type="character")
    _configure_style_fonts(style, text_sample="程式碼 コード code", is_code=True, size=Pt(9))
    font = style.font
    font.color.rgb = RGBColor(0xE3, 0x11, 0x6C)


def _ensure_quote_style(doc):
    style = _get_or_create_style(doc, "Block Quote", "Normal")
    _configure_style_fonts(style, text_sample="繁體中文 日本語 Français Español Deutsch")
    font = style.font
    font.color.rgb = RGBColor(0x6A, 0x73, 0x7D)
    font.italic = True
    pf = style.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)


def _ensure_list_styles(doc):
    for name in ["List Bullet", "List Number", "List Bullet 2", "List Number 2",
                 "List Bullet 3", "List Number 3"]:
        style = _get_or_create_style(doc, name, "Normal")
        _configure_style_fonts(style, text_sample="繁體中文 日本語 Français Español Deutsch")


def _set_paragraph_shading(pPr_element, hex_color: str):
    """Apply a background fill to a paragraph style element."""
    pPr = pPr_element.get_or_add_pPr() if hasattr(pPr_element, 'get_or_add_pPr') else pPr_element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_paragraph_shading(paragraph, hex_color: str):
    """Apply background shading to a specific paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_left_border(paragraph, hex_color: str = "DFE2E5", width_eighths: int = 24):
    """Add a left border to a paragraph (blockquote style)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_eighths))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_horizontal_rule(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D7DE")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# Publishing-quality formatting
# ---------------------------------------------------------------------------

class _HeadingConfig(TypedDict, total=False):
    size: Pt
    bold: bool
    italic: bool
    space_before: Pt
    space_after: Pt


_PUBLISH_HEADING_CONFIG: dict[int, _HeadingConfig] = {
    1: {"size": Pt(22), "bold": True,  "space_before": Pt(24), "space_after": Pt(12)},
    2: {"size": Pt(18), "bold": True,  "space_before": Pt(18), "space_after": Pt(8)},
    3: {"size": Pt(14), "bold": True,  "space_before": Pt(14), "space_after": Pt(6)},
    4: {"size": Pt(12), "bold": True,  "space_before": Pt(12), "space_after": Pt(4), "italic": True},
    5: {"size": Pt(11), "bold": False, "space_before": Pt(10), "space_after": Pt(4), "italic": True},
    6: {"size": Pt(11), "bold": False, "space_before": Pt(10), "space_after": Pt(4), "italic": True},
}


def apply_publish_formatting(doc: Document) -> None:
    """Apply publishing-quality formatting to the document.

    Enhancements applied:
    - Page margins: 1.25 in left/right, 1 in top/bottom
    - Body text: 12 pt, 1.5x line spacing, justified alignment, widow/orphan control
    - Heading hierarchy: scaled sizes, generous spacing, keep-with-next
    - Centered page-number footer on every section
    """
    _set_publish_page_margins(doc)
    _apply_publish_body_style(doc)
    _apply_publish_heading_styles(doc)
    _add_page_number_footer(doc)


def _set_publish_page_margins(doc: Document) -> None:
    """Set publication-standard page margins (1.25 in left/right, 1 in top/bottom)."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _apply_publish_body_style(doc: Document) -> None:
    """Apply publishing typography to Normal style: 12 pt, 1.5x spacing, justified."""
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.widow_control = True


def _apply_publish_heading_styles(doc: Document) -> None:
    """Apply publication-quality heading sizes, spacing, and keep-with-next."""
    for level, cfg in _PUBLISH_HEADING_CONFIG.items():
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.size = cfg["size"]
        style.font.bold = cfg["bold"]
        if cfg.get("italic"):
            style.font.italic = True
        pf = style.paragraph_format
        pf.space_before = cfg["space_before"]
        pf.space_after = cfg["space_after"]
        pf.keep_with_next = True
        pf.widow_control = True


def _add_page_number_footer(doc: Document) -> None:
    """Add a centered page-number footer to every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar_begin)
        run._element.append(instr)
        run._element.append(fldChar_end)

